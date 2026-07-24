#!/usr/bin/env python3
"""Generate ATS-safe DOCX from tailored CV markdown."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_md(md_text: str) -> list[dict]:
    """Parse markdown into structured blocks."""
    blocks = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            blocks.append({"type": "h1", "text": stripped[2:].strip()})
            i += 1
        elif stripped.startswith("## "):
            blocks.append({"type": "h2", "text": stripped[3:].strip()})
            i += 1
        elif stripped.startswith("### "):
            blocks.append({"type": "h3", "text": stripped[4:].strip()})
            i += 1
        elif stripped == "---":
            blocks.append({"type": "hr"})
            i += 1
        elif stripped.startswith("- "):
            bullet_text = stripped[2:].strip()
            sub_bullets = []
            i += 1
            while i < len(lines):
                nl = lines[i].strip()
                if nl.startswith("  - "):
                    sub_bullets.append(nl[4:].strip())
                    i += 1
                elif nl.startswith("- "):
                    break
                elif nl == "":
                    break
                else:
                    break
            blocks.append({
                "type": "bullet",
                "text": bullet_text,
                "sub_bullets": sub_bullets,
            })
        else:
            blocks.append({"type": "para", "text": stripped})
            i += 1

    return blocks


def add_inline_formatting(paragraph, text: str):
    """Add bold inline formatting from markdown **text**."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def create_docx(md_path: str, docx_path: str):
    """Create DOCX from markdown CV."""
    md_text = Path(md_path).read_text(encoding="utf-8")
    blocks = parse_md(md_text)

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Liberation Sans"
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.05

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    for block in blocks:
        if block["type"] == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(block["text"])
            run.bold = True
            run.font.size = Pt(16)
            p.space_after = Pt(2)

        elif block["type"] == "h2":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True

        elif block["type"] == "h3":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.bold = True
            run.font.size = Pt(10.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.keep_with_next = True

        elif block["type"] == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

        elif block["type"] == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            add_inline_formatting(p, block["text"])
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            for sub in block.get("sub_bullets", []):
                sp = doc.add_paragraph(style="List Bullet 2")
                add_inline_formatting(sp, sub)
                sp.paragraph_format.space_after = Pt(0)
                sp.paragraph_format.line_spacing = 1.05

        elif block["type"] == "para":
            p = doc.add_paragraph()
            add_inline_formatting(p, block["text"])
            p.paragraph_format.space_after = Pt(2)

    doc.save(docx_path)
    print(f"DOCX created: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    create_docx(sys.argv[1], sys.argv[2])
